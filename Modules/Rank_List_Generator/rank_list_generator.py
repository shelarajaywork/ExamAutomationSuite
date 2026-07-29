import io
import re
import datetime
import math
import pandas as pd
import docx
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import streamlit as st


# ==========================================
# COLLEGE NAME MAPPER
# ==========================================

COLLEGE_FULL_NAMES = {
    "Mithibai College": "Mithibai College of Arts, Chauhan Institute of Science & Amrutben Jivanlal College of Commerce and Economics (Autonomous)",
    "N. M. College": "Narsee Monjee College of Commerce and Economics (Empowered Autonomous)",
    "U.P.G. College": "Usha Pravin Gandhi College of Arts, Science and Commerce (Autonomous)"
}


# ==========================================
# HELPER FUNCTIONS
# ==========================================

def clean_rank(val):
    """Extracts integer rank from values like '1.0', 'Rank 1', or numeric 1."""
    if pd.isna(val):
        return None
    val_str = str(val).strip()
    match = re.search(r'\d+', val_str)
    if match:
        return int(match.group(0))
    return None


def find_column(df, target_name):
    """Finds matching column name regardless of minor whitespace/case differences."""
    target_clean = re.sub(r'[^a-zA-Z0-9]', '', target_name).lower()
    for col in df.columns:
        col_clean = re.sub(r'[^a-zA-Z0-9]', '', str(col)).lower()
        if target_clean == col_clean:
            return col
    return None


def parse_academic_batches(academic_year):
    """Calculates UG and PG batch ranges dynamically from the Academic Year string."""
    ay_str = str(academic_year).strip()
    match = re.search(r'(\d{4})\s*-\s*(\d{2,4})', ay_str)
    if match:
        start_yr = int(match.group(1))
        end_yr_raw = int(match.group(2))
        end_yr = (start_yr // 100) * 100 + end_yr_raw if end_yr_raw < 100 else end_yr_raw
        
        ug_start, ug_end = end_yr - 3, end_yr
        pg_start, pg_end = end_yr - 2, end_yr
        
        ug_str = f"UG (Batch of {ug_start}-{ug_end})"
        pg_str = f"PG (Batch of {pg_start}-{pg_end})"
        return ug_str, pg_str
    else:
        return "UG (Batch of 2023-2026)", "PG (Batch of 2024-2026)"


def set_cell_shading(cell, color_hex="A1A1A1"):
    """Applies background fill color to a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))


def format_cell(cell, text="", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex=None):
    """Applies Times New Roman 11pt, alignment, single line spacing, and optional fill color to table cells."""
    cell.text = str(text)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill_hex:
        set_cell_shading(cell, fill_hex)
        
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = bold


def apply_table_borders(table):
    """Applies standard full borders to the Word table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        r'<w:tblBorders %s>'
        r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        r'</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)


# ==========================================
# REPORT GENERATOR 1: TOP 5% MERITORIOUS STUDENTS
# ==========================================

def generate_word_report(df_combined, college_name, academic_year):
    """Generates the Top 5% Meritorious Toppers Word Document."""
    doc = Document()
    
    # Page Setup: Landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    
    section.top_margin = Cm(2.7)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    
    # Normal Style configuration: Line Spacing 1.5
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    
    # Page Header (Includes 1 blank line at bottom)
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.line_spacing = 1.0
    header_p.paragraph_format.space_after = Pt(0)
    
    r1 = header_p.add_run("SVKM's\n")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(11)
    r1.font.bold = True
    
    r2 = header_p.add_run(f"{college_name}\n")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    r2.font.bold = True
    r2.font.underline = True
    
    r3 = header_p.add_run("List of Rank Holders on the basis of Academic Performance\n")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(11)
    r3.font.bold = True
    r3.font.underline = True
    
    # Calculate dynamic UG & PG batches
    ug_batch_str, pg_batch_str = parse_academic_batches(academic_year)
    
    # Intro Paragraph (Line Spacing 1.5)
    intro_p = doc.add_paragraph()
    intro_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_p.paragraph_format.line_spacing = 1.5
    intro_p.paragraph_format.space_after = Pt(12)
    intro_p.paragraph_format.space_before = Pt(0)
    
    intro_text = (
        f"Following is the list of Meritorious Students (Top 5% ranks in the respective batches) for the "
        f"Academic Year {academic_year}, on the basis of academic performance of {college_name}, "
        f"based on the highest CGPA (Cumulative Grade Point Average) of their respective {ug_batch_str} "
        f"and {pg_batch_str} Programs and batches, as per the approved examination rules under autonomy."
    )
    r_intro = intro_p.add_run(intro_text)
    r_intro.font.name = "Times New Roman"
    r_intro.font.size = Pt(11)
    
    # Table Building
    table = doc.add_table(rows=0, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table)
    
    col_widths = [Cm(1.5), Cm(3.5), Cm(8.0), Cm(8.0), Cm(2.0), Cm(2.7)]
    sr_no = 1
    
    programs = df_combined['Program Name'].unique()
    
    for prog in programs:
        prog_df = df_combined[df_combined['Program Name'] == prog]
        total_students = len(prog_df)
        top_count = max(1, math.ceil(total_students * 0.05))
        
        # Banner Row (0.9 cm height)
        banner_row = table.add_row()
        banner_row.height = Cm(0.9)
        banner_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        
        banner_cell = banner_row.cells[0]
        for c in banner_row.cells[1:]:
            banner_cell = banner_cell.merge(c)
            
        banner_text = f"{prog} : No of students:- {total_students} ({top_count} ranks for meritorious students)"
        format_cell(banner_cell, banner_text, bold=True)
        
        # Column Headers (0.9 cm height, filled with #A1A1A1)
        hdr_row = table.add_row()
        hdr_row.height = Cm(0.9)
        hdr_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        
        headers = ["Sr.", "Student No", "Student Name", "Program Name", "CGPA", "Certificate"]
        for idx, h_text in enumerate(headers):
            format_cell(hdr_row.cells[idx], h_text, bold=True, fill_hex="A1A1A1")
            
        # Top 5% Student Records (0.9 cm height)
        toppers = prog_df[prog_df['CleanRank'].notna() & (prog_df['CleanRank'] <= top_count)]
        
        for _, stud in toppers.iterrows():
            row = table.add_row()
            row.height = Cm(0.9)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            
            cgpa_val = stud['CGPA']
            cgpa_str = f"{cgpa_val:.2f}" if isinstance(cgpa_val, (int, float)) else str(cgpa_val)
            
            data_vals = [
                str(sr_no),
                stud['Student No'],
                stud['Student Name'],
                stud['Program Name'],
                cgpa_str,
                ""
            ]
            for idx, val in enumerate(data_vals):
                format_cell(row.cells[idx], val, bold=False)
            sr_no += 1
            
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = width

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# ==========================================
# REPORT GENERATOR 2: FIRST RANK HOLDERS ONLY
# ==========================================

def generate_first_rankers_report(df_combined, college_name, academic_year, annual_day_date_str):
    """Generates the First Rank Holders Word Document with Endowment format & Signatures."""
    doc = Document()
    
    # Page Setup: Landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    
    section.top_margin = Cm(2.7)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    
    # Normal Style configuration: Line Spacing 1.5
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    
    # Page Header (Includes 1 blank line at bottom)
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.line_spacing = 1.0
    header_p.paragraph_format.space_after = Pt(0)
    
    r1 = header_p.add_run("SVKM's\n")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(11)
    r1.font.bold = True
    
    r2 = header_p.add_run(f"{college_name}\n")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    r2.font.bold = True
    r2.font.underline = True
    
    r3 = header_p.add_run("List of Toppers (First Rank Holders)\n")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(11)
    r3.font.bold = True
    r3.font.underline = True
    
    # Intro Paragraph (Line Spacing 1.5)
    intro_p = doc.add_paragraph()
    intro_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_p.paragraph_format.line_spacing = 1.5
    intro_p.paragraph_format.space_after = Pt(12)
    intro_p.paragraph_format.space_before = Pt(0)
    
    intro_text = (
        f"Following is the list of Toppers (First Rank Holders) for the Academic Year {academic_year}, "
        f"on the basis of academic performance of {college_name}, based on the highest CGPA of their "
        f"respective UG and PG Programs and batches, as per the approved rules. They will be felicitated "
        f"during the Prize Distribution ceremony to be held on the College Annual Day scheduled on {annual_day_date_str}."
    )
    r_intro = intro_p.add_run(intro_text)
    r_intro.font.name = "Times New Roman"
    r_intro.font.size = Pt(11)
    
    # Table Creation (6 Columns)
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table)
    
    hdr_row = table.rows[0]
    hdr_row.height = Cm(0.9)
    hdr_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    
    headers = [
        "Sr. No.",
        "Programme",
        "Name of the Student",
        "SAP No.\n(Roll No)",
        "Overall Topper\n(CGPA)",
        "Endowment Prize Amount"
    ]
    for idx, h_text in enumerate(headers):
        format_cell(hdr_row.cells[idx], h_text, bold=True, fill_hex="A1A1A1")
        
    col_widths = [Cm(1.5), Cm(7.5), Cm(7.5), Cm(3.5), Cm(2.7), Cm(3.0)]
    
    # Filter 1st Rank Holders
    rank1_df = df_combined[df_combined['CleanRank'] == 1].copy()
    
    sr_no = 1
    for _, stud in rank1_df.iterrows():
        row = table.add_row()
        row.height = Cm(0.9)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        
        cgpa_val = stud['CGPA']
        cgpa_str = f"{cgpa_val:.2f}" if isinstance(cgpa_val, (int, float)) else str(cgpa_val)
        
        data_vals = [
            str(sr_no),
            stud['Program Name'],
            stud['Student Name'],
            stud['Student No'],
            cgpa_str,
            ""
        ]
        for idx, val in enumerate(data_vals):
            format_cell(row.cells[idx], val, bold=False)
        sr_no += 1
        
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = width

    # Signatures Table Section
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    sig_row1 = sig_table.rows[0]
    format_cell(sig_row1.cells[0], "Prepared and Verified by:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    format_cell(sig_row1.cells[1], "Controller of Examinations:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    
    sig_row2 = sig_table.rows[1]
    format_cell(sig_row2.cells[0], "EEC Convener:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    format_cell(sig_row2.cells[1], "Principal:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    
    for row in sig_table.rows:
        row.cells[0].width = Cm(13.0)
        row.cells[1].width = Cm(12.7)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(10)
            cell.paragraphs[0].paragraph_format.space_after = Pt(10)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# ==========================================
# MODULE ENTRY POINT
# ==========================================

def show():
    st.title("🏆 Academic Toppers List Generator")
    st.markdown("Upload student result Excel files to automatically calculate **Top 5% Meritorious Students** and **First Rank Holders** and generate formatted Word reports.")
    st.markdown("---")
    
    # ------------------------------------------
    # CONFIGURATION INPUTS
    # ------------------------------------------
    cfg_col1, cfg_col2, cfg_col3 = st.columns([2, 1, 1])
    
    with cfg_col1:
        college_options = ["", "Mithibai College", "N. M. College", "U.P.G. College"]
        selected_college_short = st.selectbox("Select College Name", options=college_options, index=0)
        
        college_name_full = COLLEGE_FULL_NAMES.get(selected_college_short, selected_college_short)
        
    with cfg_col2:
        ay_options = [
            "",
            "2021-22",
            "2022-23",
            "2023-24",
            "2024-25",
            "2025-26",
            "2026-27",
            "2027-28",
            "2028-29",
            "2029-30"
        ]
        academic_year = st.selectbox("Academic Year", options=ay_options, index=0)
        
    with cfg_col3:
        annual_day_date = st.date_input("Date of Annual Day", value=datetime.date(2026, 1, 10), format="DD-MM-YYYY")
        # Format date as "Day, Month DD, YYYY" (e.g. "Saturday, January 10, 2026")
        annual_day_date_str = f"{annual_day_date.strftime('%A, %B')} {annual_day_date.day}, {annual_day_date.year}"

    st.markdown("---")
    
    # ------------------------------------------
    # FILE UPLOAD AREA
    # ------------------------------------------
    uploaded_files = st.file_uploader(
        "Upload Excel Files (.xlsx / .xls)", 
        type=["xlsx", "xls"], 
        accept_multiple_files=True
    )

    if uploaded_files:
        all_rows = []
        
        for file in uploaded_files:
            xls = pd.ExcelFile(file)
            for sheet in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet)
                
                col_stud_no = find_column(df, "Student No")
                col_stud_name = find_column(df, "Student Name")
                col_prog = find_column(df, "Program Name")
                col_cgpa = find_column(df, "CGPA")
                col_rank = find_column(df, "Rank")
                
                if not all([col_stud_no, col_stud_name, col_prog, col_cgpa, col_rank]):
                    continue
                    
                for _, row in df.iterrows():
                    prog_name = str(row[col_prog]).strip() if pd.notna(row[col_prog]) else ""
                    if not prog_name:
                        continue
                        
                    all_rows.append({
                        'Student No': str(row[col_stud_no]) if pd.notna(row[col_stud_no]) else "",
                        'Student Name': str(row[col_stud_name]) if pd.notna(row[col_stud_name]) else "",
                        'Program Name': prog_name,
                        'CGPA': row[col_cgpa],
                        'Rank': row[col_rank],
                        'CleanRank': clean_rank(row[col_rank])
                    })

        if all_rows:
            df_combined = pd.DataFrame(all_rows)
            df_combined.sort_values(by=['Program Name', 'CleanRank'], ascending=[True, True], na_position='last', inplace=True)
            
            st.success(f"Successfully processed **{len(uploaded_files)}** file(s) across **{df_combined['Program Name'].nunique()}** programs.")
            
            # Summary Table (Renamed Columns)
            summary_data = []
            for prog in df_combined['Program Name'].unique():
                prog_df = df_combined[df_combined['Program Name'] == prog]
                tot = len(prog_df)
                top = max(1, round(tot * 0.05))
                summary_data.append({
                    "Program Name": prog,
                    "Total Students": tot,
                    "Top 5% Count": top
                })
            
            st.subheader("📊 Program Summary Preview")
            st.dataframe(pd.DataFrame(summary_data), use_container_width=True)
            
            if not selected_college_short or not academic_year:
                st.warning("⚠️ Please select a College Name and Academic Year from the dropdowns above before generating the Word documents.")
            
            # Generate Word Documents
            if st.button("Generate Word Documents", type="primary", disabled=(not selected_college_short or not academic_year)):
                with st.spinner("Generating formatted Word documents..."):
                    top5_doc = generate_word_report(df_combined, college_name_full, academic_year)
                    rank1_doc = generate_first_rankers_report(df_combined, college_name_full, academic_year, annual_day_date_str)
                    
                    st.success("✅ Both Word documents generated successfully! Click below to download:")
                    
                    btn_col1, btn_col2 = st.columns(2)
                    
                    with btn_col1:
                        st.download_button(
                            label="📥 Download Top 5% Rank List (Topper_List.docx)",
                            data=top5_doc,
                            file_name="Topper_List_Top5Percent.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        
                    with btn_col2:
                        st.download_button(
                            label="📥 Download 1st Rank Holders List (First_Rank_Holders.docx)",
                            data=rank1_doc,
                            file_name="First_Rank_Holders.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
        else:
            st.error("Could not find required columns (Student No, Student Name, Program Name, CGPA, Rank) in uploaded files.")


# Enable standalone execution when run directly via Streamlit
if __name__ == "__main__":
    st.set_page_config(page_title="Rank List Generator", layout="wide")
    show()